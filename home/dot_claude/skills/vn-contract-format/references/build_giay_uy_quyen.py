"""Generate print-ready DOCX of Giấy uỷ quyền - Trang -> Nhiên."""
from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/tieubao/workspace/tieubao/properties/listings/le-hong-phong-dalat-house/contracts/2026-05-06_handover/04-giay-uy-quyen.docx"

doc = Document()
for section in doc.sections:
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts')
    rPr.append(rFonts)
rFonts.set(qn('w:ascii'), 'Times New Roman')
rFonts.set(qn('w:hAnsi'), 'Times New Roman')
rFonts.set(qn('w:cs'), 'Times New Roman')

def add_para(text, *, bold=False, italic=False, align=None, size=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p

def add_blank(size=Pt(6)):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = size
    p.paragraph_format.space_before = Pt(0)

def borderless(table):
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for edge in ('top', 'left', 'bottom', 'right'):
                b = OxmlElement(f'w:{edge}')
                b.set(qn('w:val'), 'nil')
                tcBorders.append(b)
            tcPr.append(tcBorders)

# Header
add_para("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
add_para("Độc lập - Tự do - Hạnh phúc", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
add_para("―――――― o0o ――――――", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(18))

# Title
add_para("GIẤY UỶ QUYỀN", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=18, space_after=Pt(2))
add_para("(V/v: đại diện ký Biên bản thanh lý hợp đồng thuê nhà)", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_after=Pt(16))

# Intro
intro = doc.add_paragraph()
intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
intro.paragraph_format.space_after = Pt(10)
intro.add_run("Hôm nay, ngày ___ tháng ___ năm 2026, tại Thành phố ____________, chúng tôi gồm:")

def set_fixed_table_layout(table, col_widths_cm):
    tblPr = table._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    for el in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(el)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    tblGrid = table._tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        table._tbl.remove(tblGrid)
    tblGrid = OxmlElement('w:tblGrid')
    for w_cm in col_widths_cm:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(int(w_cm * 567)))
        tblGrid.append(gridCol)
    tblPr.addnext(tblGrid)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:w'), str(int(col_widths_cm[i] * 567)))
            tcW.set(qn('w:type'), 'dxa')


def apply_borders(table, fillin_mask):
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            for el in tcPr.findall(qn('w:tcBorders')):
                tcPr.remove(el)
            tcBorders = OxmlElement('w:tcBorders')
            for edge in ('top', 'left', 'right'):
                b = OxmlElement(f'w:{edge}')
                b.set(qn('w:val'), 'nil')
                tcBorders.append(b)
            bottom = OxmlElement('w:bottom')
            if fillin_mask[i][j]:
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '6')
                bottom.set(qn('w:color'), '000000')
            else:
                bottom.set(qn('w:val'), 'nil')
            tcBorders.append(bottom)
            tcPr.append(tcBorders)


def info_table(rows, label_w_cm=4.5, value_w_cm=12.2):
    t = doc.add_table(rows=len(rows), cols=2)
    t.autofit = False
    set_fixed_table_layout(t, [label_w_cm, value_w_cm])
    fillin_mask = []
    for i, (label, value) in enumerate(rows):
        c0, c1 = t.rows[i].cells
        c0.paragraphs[0].text = ""
        run = c0.paragraphs[0].add_run(label)
        run.font.size = Pt(13)
        c1.paragraphs[0].text = ""
        is_fillin = value.startswith("____") or value == ""
        if not is_fillin:
            run = c1.paragraphs[0].add_run(value)
            run.font.size = Pt(13)
            if value.startswith("Bà NGUYỄN") or value.startswith("Bà TRẦN"):
                run.bold = True
        fillin_mask.append([False, is_fillin])
    apply_borders(t, fillin_mask)
    return t

# Bên A
add_para("BÊN UỶ QUYỀN (BÊN A)", bold=True, space_after=Pt(4))
info_table([
    ("Họ và tên:", "Bà NGUYỄN THỊ MINH TRANG"),
    ("Năm sinh:", "____________________________________"),
    ("CCCD/CMND số:", "____________________________________"),
    ("Ngày cấp:", "____________________________________"),
    ("Nơi cấp:", "____________________________________"),
    ("Địa chỉ thường trú:", "____________________________________"),
    ("Số điện thoại:", "____________________________________"),
])
add_blank(Pt(8))

# Bên B
add_para("BÊN ĐƯỢC UỶ QUYỀN (BÊN B)", bold=True, space_after=Pt(4))
info_table([
    ("Họ và tên:", "Bà TRẦN HOÀNG HẠO NHIÊN"),
    ("Năm sinh:", "____________________________________"),
    ("CCCD/CMND số:", "____________________________________"),
    ("Ngày cấp:", "____________________________________"),
    ("Nơi cấp:", "____________________________________"),
    ("Địa chỉ thường trú:", "____________________________________"),
    ("Số điện thoại:", "____________________________________"),
])
add_blank(Pt(10))

# Body intro
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(10)
p.add_run("Bằng văn bản này, Bên A đồng ý uỷ quyền cho Bên B thực hiện các nội dung sau:")

# Điều 1
add_para("ĐIỀU 1. NỘI DUNG UỶ QUYỀN", bold=True, space_after=Pt(4))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(6)
p.add_run("Bên A uỷ quyền cho Bên B thay mặt Bên A thực hiện đầy đủ các công việc sau đây liên quan đến việc bàn giao và thanh lý hợp đồng thuê nhà tại địa chỉ ")
r = p.add_run("5/5 Lê Hồng Phong, Phường Xuân Hương, Thành phố Đà Lạt"); r.bold = True
p.add_run(":")

for i, line in enumerate([
    "Kiểm tra hiện trạng nhà, tài sản, thiết bị tại thời điểm bên thuê trả nhà.",
    "Chốt số đồng hồ điện (mã KH PB03010012398), đồng hồ nước (mã KH 106172).",
    "Tiếp nhận thanh toán các khoản công nợ (điện, nước, internet, rác) từ bên thuê.",
    "Thu hồi toàn bộ chìa khoá nhà.",
    "Đại diện Bên cho thuê ký Biên bản thanh lý hợp đồng thuê nhà với bên thuê.",
    "Thực hiện các công việc khác có liên quan để hoàn tất việc bàn giao trả nhà.",
], start=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.7)
    p.add_run(f"{i}. {line}").font.size = Pt(13)

add_blank(Pt(8))

# Điều 2
add_para("ĐIỀU 2. PHẠM VI VÀ THỜI HẠN UỶ QUYỀN", bold=True, space_after=Pt(4))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(6)
r = p.add_run("2.1. "); r.bold = True
p.add_run("Phạm vi uỷ quyền: chỉ giới hạn trong các công việc nêu tại Điều 1 của Giấy uỷ quyền này. Bên B không được uỷ quyền lại cho người khác và không được thực hiện bất kỳ giao dịch nào ngoài phạm vi nêu trên.")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(10)
r = p.add_run("2.2. "); r.bold = True
p.add_run("Thời hạn uỷ quyền: từ ngày ký Giấy uỷ quyền này cho đến khi hoàn tất việc bàn giao và ký Biên bản thanh lý hợp đồng thuê nhà (dự kiến ngày 06/05/2026).")

# Điều 3
add_para("ĐIỀU 3. QUYỀN VÀ NGHĨA VỤ CỦA CÁC BÊN", bold=True, space_after=Pt(4))
for label, body in [
    ("3.1.", "Bên A có trách nhiệm cung cấp đầy đủ thông tin, tài liệu cần thiết để Bên B thực hiện công việc được uỷ quyền (bao gồm hợp đồng thuê nhà, biên bản bàn giao nhận nhà ban đầu, danh mục tài sản, thông tin tài khoản nhận thanh toán)."),
    ("3.2.", "Bên B có trách nhiệm thực hiện công việc được uỷ quyền một cách trung thực, đầy đủ và đúng phạm vi; bàn giao lại toàn bộ tài liệu, chìa khoá và báo cáo kết quả cho Bên A ngay sau khi hoàn tất."),
    ("3.3.", "Bên A chịu trách nhiệm pháp lý đối với các hành vi của Bên B trong phạm vi uỷ quyền."),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"{label} "); r.bold = True
    p.add_run(body)
add_blank(Pt(4))

# Điều 4
add_para("ĐIỀU 4. CAM KẾT CHUNG", bold=True, space_after=Pt(4))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(10)
p.add_run("Hai bên cam kết thực hiện đúng và đầy đủ các nội dung tại Giấy uỷ quyền này. Mọi tranh chấp (nếu có) sẽ được giải quyết trên cơ sở thoả thuận; trường hợp không thoả thuận được sẽ giải quyết theo quy định pháp luật hiện hành.")

# Điều 5
add_para("ĐIỀU 5. HIỆU LỰC", bold=True, space_after=Pt(4))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(16)
p.add_run("Giấy uỷ quyền này có hiệu lực kể từ ngày ký và được lập thành ")
r = p.add_run("02 (hai) bản"); r.bold = True
p.add_run(" có giá trị pháp lý như nhau, mỗi bên giữ ")
r = p.add_run("01 (một) bản"); r.bold = True
p.add_run(".")

# Signature table
sig_t = doc.add_table(rows=4, cols=2)
sig_t.autofit = False

def set_cell(cell, text, *, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=13):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)

set_cell(sig_t.rows[0].cells[0], "BÊN UỶ QUYỀN (BÊN A)", bold=True)
set_cell(sig_t.rows[0].cells[1], "BÊN ĐƯỢC UỶ QUYỀN (BÊN B)", bold=True)
set_cell(sig_t.rows[1].cells[0], "(Ký, ghi rõ họ tên)", italic=True, size=11)
set_cell(sig_t.rows[1].cells[1], "(Ký, ghi rõ họ tên)", italic=True, size=11)
set_cell(sig_t.rows[2].cells[0], "")
set_cell(sig_t.rows[2].cells[1], "")
set_cell(sig_t.rows[3].cells[0], "NGUYỄN THỊ MINH TRANG", bold=True)
set_cell(sig_t.rows[3].cells[1], "TRẦN HOÀNG HẠO NHIÊN", bold=True)

# Add space below signature row 2 for handwritten signature
sig_t.rows[2].height = Cm(2.2)

borderless(sig_t)
add_blank(Pt(14))

# Đính kèm
add_para("Tài liệu đính kèm:", bold=True, space_after=Pt(4))
for line in [
    "□ Bản sao CCCD/CMND của Bên uỷ quyền (Bà Nguyễn Thị Minh Trang)",
    "□ Bản sao CCCD/CMND của Bên được uỷ quyền (Bà Trần Hoàng Hạo Nhiên)",
]:
    p = doc.add_paragraph(line)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)

doc.save(OUT)
print(f"OK: {OUT}")
