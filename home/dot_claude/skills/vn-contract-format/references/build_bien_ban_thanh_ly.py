"""Generate print-ready DOCX of Biên bản thanh lý hợp đồng thuê nhà."""
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/tieubao/workspace/tieubao/properties/listings/le-hong-phong-dalat-house/contracts/2026-05-06_handover/03-bien-ban-thanh-ly.docx"

doc = Document()

# A4 + 2cm margins
for section in doc.sections:
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2)

# Default font: Times New Roman 13 (Vietnamese standard)
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts')
    rPr.append(rFonts)
rFonts.set(qn('w:ascii'), 'Times New Roman')
rFonts.set(qn('w:hAnsi'), 'Times New Roman')
rFonts.set(qn('w:cs'), 'Times New Roman')

def add_para(text, *, bold=False, italic=False, align=None, size=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p

def add_blank(size=Pt(6)):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = size
    p.paragraph_format.space_before = Pt(0)


def add_label_with_leader_line(cell, label, *, tab_at_cm=7.5, font_size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(2)):
    """Cell content: 'Label' + tab character. Tab stop at tab_at_cm with underscore-line leader.
    Word draws an underscore line from end of label to tab stop position - auto-fits, never wraps."""
    cell.paragraphs[0].text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(
        Cm(tab_at_cm), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.LINES
    )
    if label:
        run = p.add_run(label)
        run.font.size = Pt(font_size)
    run = p.add_run("\t")
    run.font.size = Pt(font_size)
    return p


def fill_lines_table(n_rows, *, width_cm=16.7, row_height_cm=0.7, parent=None, space_after=Pt(6)):
    """N-row, single-column table with bottom border per cell. Use for full-width fill-in lines."""
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    if parent is None:
        parent = doc
    t = parent.add_table(rows=n_rows, cols=1)
    t.autofit = False
    set_fixed_table_layout(t, [width_cm])
    for i in range(n_rows):
        row = t.rows[i]
        row.height = Cm(row_height_cm)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        cell = row.cells[0]
        cell.paragraphs[0].text = ""
    # Borders: bottom on each cell, others nil
    for row in t.rows:
        cell = row.cells[0]
        tcPr = cell._tc.get_or_add_tcPr()
        for el in tcPr.findall(qn('w:tcBorders')):
            tcPr.remove(el)
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'right'):
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'), 'nil')
            tcBorders.append(b)
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:color'), '000000')
        tcBorders.append(bottom)
        tcPr.append(tcBorders)
    # Add small spacer paragraph after table to give breathing room
    if parent is doc:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = space_after
        sp.paragraph_format.space_before = Pt(0)
    return t


def add_fill_line(label="", *, indent_cm=0, font_size=13, space_after=Pt(8), italic_label=False):
    """Paragraph with optional 'Label:' followed by a horizontal line (paragraph bottom border)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    if label:
        run = p.add_run(label)
        run.font.size = Pt(font_size)
        run.italic = italic_label
    _set_para_bottom_border(p)
    return p


def add_fill_line_in_cell(cell, label="", *, font_size=13, space_after=Pt(4)):
    """Paragraph inside a table cell with 'Label:' + bottom border line.
    Reuses the first empty paragraph if present, else appends a new one."""
    first_p = cell.paragraphs[0]
    if not first_p.runs and first_p.text == "":
        p = first_p
    else:
        p = cell.add_paragraph()
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    if label:
        run = p.add_run(label)
        run.font.size = Pt(font_size)
    _set_para_bottom_border(p)
    return p

# Header: quốc hiệu + tiêu ngữ
add_para("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
add_para("Độc lập - Tự do - Hạnh phúc", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
add_para("―――――― o0o ――――――", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(18))

# Title
add_para("BIÊN BẢN THANH LÝ HỢP ĐỒNG THUÊ NHÀ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, space_after=Pt(14))

# Căn cứ
add_para("- Căn cứ Hợp đồng thuê nhà ký ngày ___/___/______ giữa Bên cho thuê và Bên thuê;", italic=True, space_after=Pt(2))
add_para("- Căn cứ nguyện vọng tự nguyện chấm dứt Hợp đồng thuê nhà trước thời hạn của Bên thuê.", italic=True, space_after=Pt(10))

# Intro
intro = doc.add_paragraph()
intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
intro.paragraph_format.space_after = Pt(10)
r1 = intro.add_run("Hôm nay, ngày ")
r2 = intro.add_run("06 tháng 05 năm 2026"); r2.bold = True
intro.add_run(", tại địa chỉ ")
r3 = intro.add_run("5/5 Lê Hồng Phong, Phường Xuân Hương, Thành phố Đà Lạt"); r3.bold = True
intro.add_run(", chúng tôi gồm:")

# Bên A
add_para("BÊN CHO THUÊ (BÊN A)", bold=True, space_after=Pt(4))

def set_fixed_table_layout(table, col_widths_cm):
    """Force fixed column widths via tblLayout + tblGrid + tcW."""
    tblPr = table._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    for el in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(el)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    tblGrid = table._tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        table._tbl.remove(tblGrid)
    tblGrid = OxmlElement('w:tblGrid')
    for w_cm in col_widths_cm:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(int(w_cm * 567)))
        tblGrid.append(gridCol)
    tblPr.addnext(tblGrid)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:w'), str(int(col_widths_cm[i] * 567)))
            tcW.set(qn('w:type'), 'dxa')


def apply_borders(table, fillin_mask):
    """fillin_mask[i][j] = True means that cell gets a bottom border (fill-in line);
    all other borders are nil."""
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            for el in tcPr.findall(qn('w:tcBorders')):
                tcPr.remove(el)
            tcBorders = OxmlElement('w:tcBorders')
            for edge in ('top', 'left', 'right'):
                b = OxmlElement(f'w:{edge}')
                b.set(qn('w:val'), 'nil')
                tcBorders.append(b)
            bottom = OxmlElement('w:bottom')
            if fillin_mask[i][j]:
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '6')
                bottom.set(qn('w:color'), '000000')
            else:
                bottom.set(qn('w:val'), 'nil')
            tcBorders.append(bottom)
            tcPr.append(tcBorders)


def info_table(rows, label_w_cm=4.5, value_w_cm=12.2, *, parent=None, font_size=13, row_height_cm=None):
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    if parent is None:
        parent = doc
    t = parent.add_table(rows=len(rows), cols=2)
    t.autofit = False
    set_fixed_table_layout(t, [label_w_cm, value_w_cm])
    fillin_mask = []
    for i, (label, value) in enumerate(rows):
        if row_height_cm:
            t.rows[i].height = Cm(row_height_cm)
            t.rows[i].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        c0, c1 = t.rows[i].cells
        c0.paragraphs[0].text = ""
        run = c0.paragraphs[0].add_run(label)
        run.font.size = Pt(font_size)
        c1.paragraphs[0].text = ""
        is_fillin = value.startswith("____") or value == ""
        if not is_fillin:
            run = c1.paragraphs[0].add_run(value)
            run.font.size = Pt(font_size)
            if "NGUYỄN THỊ MINH TRANG" in value or "TRẦN HOÀNG HẠO NHIÊN" in value:
                run.bold = True
        fillin_mask.append([False, is_fillin])
    apply_borders(t, fillin_mask)
    return t

info_table([
    ("Họ và tên:", "Bà NGUYỄN THỊ MINH TRANG"),
    ("Người đại diện ký:", "Bà TRẦN HOÀNG HẠO NHIÊN (theo Giấy uỷ quyền đính kèm)"),
    ("CCCD/CMND số:", ""),
    ("Địa chỉ thường trú:", ""),
    ("Số điện thoại:", ""),
])
add_blank(Pt(8))

# Bên B
add_para("BÊN THUÊ (BÊN B)", bold=True, space_after=Pt(4))
info_table([
    ("Họ và tên (Người 1):", "____________________________________"),
    ("CCCD/CMND số:", "____________________________________"),
    ("Họ và tên (Người 2):", "____________________________________"),
    ("CCCD/CMND số:", "____________________________________"),
    ("Địa chỉ thường trú:", "____________________________________"),
    ("Số điện thoại:", "____________________________________"),
])
add_blank(Pt(10))

# Body intro
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(10)
p.add_run("Hai bên cùng thống nhất thanh lý ")
r = p.add_run("Hợp đồng thuê nhà"); r.bold = True
p.add_run(" đã ký kết với các nội dung sau:")

# Điều 1
add_para("ĐIỀU 1. LÝ DO CHẤM DỨT HỢP ĐỒNG", bold=True, space_after=Pt(4))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(10)
p.add_run("Bên B ")
r = p.add_run("tự nguyện"); r.bold = True
p.add_run(" chấm dứt Hợp đồng thuê nhà trước thời hạn vì ")
r = p.add_run("lý do cá nhân"); r.bold = True
p.add_run(", hoàn toàn ")
r = p.add_run("không do lỗi của Bên A"); r.bold = True
p.add_run(".")

# Điều 2
add_para("ĐIỀU 2. TIỀN ĐẶT CỌC", bold=True, space_after=Pt(4))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(6)
p.add_run("Bên B đồng ý ")
r = p.add_run("mất toàn bộ số tiền đặt cọc"); r.bold = True
p.add_run(" là ")
r = p.add_run("____________________ VND"); r.bold = True
p.add_run(" (bằng chữ: ___________________________________________________ đồng) theo đúng thoả thuận tại Hợp đồng thuê nhà và quy định pháp luật hiện hành.")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(10)
p.add_run("Bên A có toàn quyền sử dụng số tiền đặt cọc nói trên. Bên B không có bất kỳ quyền yêu cầu hoàn trả nào.")

# Điều 3
add_para("ĐIỀU 3. BÀN GIAO TÀI SẢN VÀ THANH TOÁN CÔNG NỢ", bold=True, space_after=Pt(4))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(4)
r = p.add_run("3.1. "); r.bold = True
p.add_run("Bên B đã bàn giao đầy đủ tại thời điểm ký Biên bản này:")

for line in [
    "Toàn bộ chìa khoá nhà, tổng cộng: _______ cái (đã được kiểm đếm và xác nhận).",
    "Toàn bộ tài sản, thiết bị thuộc sở hữu Bên A theo Biên bản bàn giao nhà ngày ___/___/_____ ; tình trạng:  □ Đầy đủ, nguyên vẹn    □ Có hư hỏng/thiếu (chi tiết tại phần Ghi chú).",
]:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.7)
    p.add_run(line).font.size = Pt(13)

add_blank(Pt(4))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(4)
r = p.add_run("3.2. "); r.bold = True
p.add_run("Bên B đã thanh toán đầy đủ các khoản công nợ tại thời điểm ký Biên bản này:")

for line in [
    "Tiền điện kỳ cuối, chỉ số đồng hồ: _______ kWh  (mã khách hàng: PB03010012398)",
    "Tiền nước kỳ cuối, chỉ số đồng hồ: _______ m³   (mã khách hàng: 106172)",
    "Tiền internet (nếu HĐ đứng tên Bên A):   □ Đã thanh toán    □ Không áp dụng",
    "Tiền rác:   □ Đã thanh toán    □ Không áp dụng",
]:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.7)
    p.add_run(line).font.size = Pt(13)

add_blank(Pt(8))

# Điều 4
add_para("ĐIỀU 4. CAM KẾT KHÔNG KHIẾU NẠI", bold=True, space_after=Pt(4))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(10)
p.add_run("Bên B cam kết đã ")
r = p.add_run("hoàn tất toàn bộ nghĩa vụ"); r.bold = True
p.add_run(" liên quan đến Hợp đồng thuê nhà và ")
r = p.add_run("sẽ không có bất kỳ khiếu nại, tranh chấp hay yêu cầu hoàn trả nào"); r.bold = True
p.add_run(" về sau, dưới bất kỳ hình thức nào, đối với Bên A.")

# Điều 5
add_para("ĐIỀU 5. HIỆU LỰC", bold=True, space_after=Pt(4))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(6)
p.add_run("Biên bản này có hiệu lực kể từ ngày ký. Hợp đồng thuê nhà giữa hai bên được ")
r = p.add_run("chính thức chấm dứt và thanh lý"); r.bold = True
p.add_run(" kể từ thời điểm này.")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(14)
p.add_run("Biên bản được lập thành ")
r = p.add_run("02 (hai) bản"); r.bold = True
p.add_run(" có giá trị pháp lý như nhau, mỗi bên giữ ")
r = p.add_run("01 (một) bản"); r.bold = True
p.add_run(".")

# Ghi chú
add_para("Ghi chú (nếu có):", bold=True, italic=True, space_after=Pt(4))
fill_lines_table(3, width_cm=16.7, row_height_cm=0.7)

add_blank(Pt(14))

# Signature table - traditional sig block layout:
#   Row 0:  BÊN A header              | BÊN B header
#   Row 1:  italic note A              | italic note B
#   Row 2:  empty (handwritten sig)    | empty (handwritten sig)
#   Row 3:  empty (handwritten sig)    | empty (handwritten sig)
#   Row 4:  TRẦN HOÀNG HẠO NHIÊN       | Người 1: ___________________
#   Row 5:  (empty)                    | Người 2: ___________________
#   Row 6:  Ngày: 06/05/2026           | Ngày: 06/05/2026
sig_t = doc.add_table(rows=7, cols=2)
sig_t.autofit = False
set_fixed_table_layout(sig_t, [8.35, 8.35])

def set_cell(cell, text, *, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=13):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)

set_cell(sig_t.rows[0].cells[0], "BÊN A (BÊN CHO THUÊ)", bold=True)
set_cell(sig_t.rows[0].cells[1], "BÊN B (BÊN THUÊ)", bold=True)
set_cell(sig_t.rows[1].cells[0], "(đại diện ký theo uỷ quyền)", italic=True, size=11)
set_cell(sig_t.rows[1].cells[1], "(cả hai vợ chồng cùng ký, ghi rõ họ tên)", italic=True, size=11)
# Rows 2-3: handwritten signature space (min height ~1.2cm each = ~2.4cm total)
for r in (2, 3):
    set_cell(sig_t.rows[r].cells[0], "")
    set_cell(sig_t.rows[r].cells[1], "")
    sig_t.rows[r].height = Cm(1.2)
    sig_t.rows[r].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
# Row 4: A=printed name, B=Người 1 + leader line
set_cell(sig_t.rows[4].cells[0], "TRẦN HOÀNG HẠO NHIÊN", bold=True)
add_label_with_leader_line(sig_t.rows[4].cells[1], "Người 1: ", tab_at_cm=7.8, font_size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(4))
# Row 5: A empty, B=Người 2 + leader line
set_cell(sig_t.rows[5].cells[0], "")
add_label_with_leader_line(sig_t.rows[5].cells[1], "Người 2: ", tab_at_cm=7.8, font_size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(4))
# Row 6: dates
set_cell(sig_t.rows[6].cells[0], "Ngày: 06/05/2026", italic=True, size=11)
set_cell(sig_t.rows[6].cells[1], "Ngày: 06/05/2026", italic=True, size=11)

# Remove borders on signature table
for row in sig_t.rows:
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'), 'nil')
            tcBorders.append(b)
        tcPr.append(tcBorders)

add_blank(Pt(14))

# Đính kèm
add_para("Tài liệu đính kèm:", bold=True, space_after=Pt(4))
for line in [
    "□ Bản sao CCCD/CMND của Bên thuê (mặt trước + mặt sau)",
    "□ Bản chụp đồng hồ điện cuối kỳ (mã PB03010012398)",
    "□ Bản chụp đồng hồ nước cuối kỳ (mã 106172)",
    "□ Chứng từ thanh toán công nợ (chuyển khoản hoặc biên nhận tiền mặt)",
    "□ Giấy uỷ quyền của Bà Nguyễn Thị Minh Trang cho Bà Trần Hoàng Hạo Nhiên (làm trước ngày ký)",
]:
    p = doc.add_paragraph(line)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)

doc.save(OUT)
print(f"OK: {OUT}")
